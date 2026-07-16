#!/usr/bin/env python3
"""Convert manuscript.md to DOCX."""
import re, os
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_tc(doc, text, style=None):
    """Add text with bold/italic formatting."""
    p = doc.add_paragraph(style=style)
    for chunk in re.split(r'(\*\*.*?\*\*|\*.*?\*)', text):
        if not chunk: continue
        if chunk.startswith('**') and chunk.endswith('**'):
            r = p.add_run(chunk[2:-2]); r.bold = True
        elif chunk.startswith('*') and chunk.endswith('*') and len(chunk)>1:
            r = p.add_run(chunk[1:-1]); r.italic = True
        else:
            p.add_run(chunk)
    return p

def convert():
    base = '/Users/rainxu/Documents/New project/paper_q1'
    with open(os.path.join(base, 'manuscript.md')) as f: lines = f.read().split('\n')
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(11)
    
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip(): i+=1; continue
        if line.startswith('# ') and not line.startswith('## '):
            add_tc(doc, line[2:], 'Title'); i+=1; continue
        if line.startswith('## '):
            add_tc(doc, line[3:], 'Heading 1'); i+=1; continue
        if line.startswith('### '):
            add_tc(doc, line[4:], 'Heading 2'); i+=1; continue
        if line.startswith('#### '):
            add_tc(doc, line[5:], 'Heading 3'); i+=1; continue
        if '|' in line and i+1 < len(lines) and '---' in lines[i+1]:
            hdr = [c.strip() for c in lines[i].split('|')[1:-1]]
            rows = []; i += 2
            while i < len(lines) and '|' in lines[i]:
                rows.append([c.strip() for c in lines[i].split('|')[1:-1]]); i+=1
            if hdr:
                t = doc.add_table(rows=1+len(rows), cols=len(hdr))
                t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j,h in enumerate(hdr):
                    r=t.rows[0].cells[j].paragraphs[0].add_run(h); r.bold=True
                for ri,row in enumerate(rows):
                    for j,val in enumerate(row):
                        if j<len(hdr): t.rows[ri+1].cells[j].text=val
            doc.add_paragraph(); continue
        add_tc(doc, line); i+=1
    
    out = os.path.join(base, 'DAFN_Q1_ready.docx')
    doc.save(out)
    print(f"DONE: {out}")

if __name__ == '__main__': convert()
